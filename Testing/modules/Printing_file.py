import docx
import fitz  # PyMuPDF
import logging
from typing import Union
import configparser
from config import validate_file_path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def load_document(file_path: str) -> Union[docx.Document, fitz.Document]:
    """
    Load a DOCX or PDF document.
    """
    try:
        if file_path.endswith('.docx'):
            return docx.Document(file_path)
        elif file_path.endswith('.pdf'):
            return fitz.open(file_path)
        else:
            raise ValueError("Unsupported file format. Only DOCX and PDF are supported.")
    except Exception as e:
        logging.error(f"Error loading document: {e}")
        raise

def modify_document(doc: Union[docx.Document, fitz.Document], new_text: str):
    """
    Modify the document by adding new text.
    """
    try:
        if isinstance(doc, docx.Document):
            doc.add_paragraph(new_text)
        elif isinstance(doc, fitz.Document):
            for page in doc:
                page.insert_text((72, 72), new_text)  # Adding text at (72, 72) coordinates
        else:
            raise ValueError("Unsupported document type.")
    except Exception as e:
        logging.error(f"Error modifying document: {e}")
        raise

def save_document(doc: Union[docx.Document, fitz.Document], new_file_path: str):
    """
    Save the modified document to a new file.
    """
    try:
        if isinstance(doc, docx.Document):
            doc.save(new_file_path)
        elif isinstance(doc, fitz.Document):
            doc.save(new_file_path)
        else:
            raise ValueError("Unsupported document type.")
    except Exception as e:
        logging.error(f"Error saving document: {e}")
        raise

def print_document(doc: Union[docx.Document, fitz.Document]):
    """
    Print the document contents.
    """
    try:
        if isinstance(doc, docx.Document):
            for paragraph in doc.paragraphs:
                print(paragraph.text)
        elif isinstance(doc, fitz.Document):
            for page in doc:
                print(page.get_text())
        else:
            raise ValueError("Unsupported document type.")
    except Exception as e:
        logging.error(f"Error printing document: {e}")
        raise

config = configparser.ConfigParser()
config_file_path = r'Prod\config.ini' # relative path
config.read(config_file_path)

save = config['paths']['SAVE']
validate_file_path(save)

input = config['paths']['INPUT']
validate_file_path(input)

def main():
    file_path = 'your_document.docx'  # or 'your_document.pdf'
    new_file_path = 'modified_document.docx'  # or 'modified_document.pdf'
    new_text = "This is the new text to be added."

    try:
        # Load the document
        doc = load_document(file_path)
        logging.info(f"Loaded document: {file_path}")

        # Modify the document
        modify_document(doc, new_text)
        logging.info("Document modified successfully.")

        # Save the document to a new file
        save_document(doc, new_file_path)
        logging.info(f"Document saved to new file: {new_file_path}")

        # Print the document
        print_document(doc)
        logging.info("Document printed successfully.")
    except Exception as e:
        logging.error(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
